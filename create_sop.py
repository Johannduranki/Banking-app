from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path

ROOT = Path(r"C:\Users\johan\Documents\Codex\2026-08-02\c")
OUT = ROOT / "Duranki_Banking_App_SOP.docx"
LOGO = ROOT / "public" / "duranki-logo.png"

NAVY = "0B1D33"; BLUE = "173A5E"; CYAN = "5BC8ED"; PALE = "E8F4F8"; GRAY = "667085"; LIGHT = "F2F4F7"; RED = "9B1C1C"; GOLD = "8A6500"; WHITE = "FFFFFF"

doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(8.5), Inches(11)
sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
sec.header_distance = sec.footer_distance = Inches(.492)

def font(run, size=11, bold=False, color="000000", name="Calibri", italic=False):
    run.font.name = name; run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name); run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size); run.bold = bold; run.italic = italic; run.font.color.rgb = RGBColor.from_string(color)
    return run

styles = doc.styles
normal = styles["Normal"]; normal.font.name="Calibri"; normal.font.size=Pt(11); normal.font.color.rgb=RGBColor.from_string("202733")
normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.25
for name,size,color,before,after in [("Title",30,NAVY,0,8),("Subtitle",14,GRAY,0,12),("Heading 1",16,BLUE,18,10),("Heading 2",13,BLUE,14,7),("Heading 3",12,"1F4D78",10,5)]:
    st=styles[name]; st.font.name="Calibri"; st.font.size=Pt(size); st.font.bold=name!="Subtitle"; st.font.color.rgb=RGBColor.from_string(color)
    st.paragraph_format.space_before=Pt(before); st.paragraph_format.space_after=Pt(after); st.paragraph_format.keep_with_next=True
for name in ["List Bullet","List Number"]:
    st=styles[name]; st.font.name="Calibri"; st.font.size=Pt(11); st.paragraph_format.left_indent=Inches(.375); st.paragraph_format.first_line_indent=Inches(-.188); st.paragraph_format.space_after=Pt(4); st.paragraph_format.line_spacing=1.25

note = styles.add_style("SOP Note", WD_STYLE_TYPE.PARAGRAPH); note.font.name="Calibri"; note.font.size=Pt(10.5); note.font.color.rgb=RGBColor.from_string(NAVY); note.paragraph_format.space_before=Pt(6); note.paragraph_format.space_after=Pt(8); note.paragraph_format.left_indent=Inches(.18); note.paragraph_format.right_indent=Inches(.18)

def shade(cell, fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=tcPr.find(qn("w:shd")) or OxmlElement("w:shd"); shd.set(qn("w:fill"),fill); tcPr.append(shd) if shd.getparent() is None else None

def margins(cell, top=80, start=120, bottom=80, end=120):
    tcPr=cell._tc.get_or_add_tcPr(); tcMar=tcPr.first_child_found_in("w:tcMar")
    if tcMar is None: tcMar=OxmlElement("w:tcMar"); tcPr.append(tcMar)
    for side,val in [("top",top),("start",start),("bottom",bottom),("end",end)]:
        el=tcMar.find(qn(f"w:{side}")) or OxmlElement(f"w:{side}"); el.set(qn("w:w"),str(val)); el.set(qn("w:type"),"dxa"); tcMar.append(el) if el.getparent() is None else None

def set_table_geometry(table, widths):
    table.autofit=False; table.alignment=WD_TABLE_ALIGNMENT.LEFT
    tblPr=table._tbl.tblPr; tblW=tblPr.find(qn("w:tblW")); tblW.set(qn("w:w"),str(sum(widths))); tblW.set(qn("w:type"),"dxa")
    tblInd=tblPr.find(qn("w:tblInd")) or OxmlElement("w:tblInd"); tblInd.set(qn("w:w"),"120"); tblInd.set(qn("w:type"),"dxa"); tblPr.append(tblInd) if tblInd.getparent() is None else None
    grid=table._tbl.tblGrid
    for old in list(grid): grid.remove(old)
    for width in widths:
        col=OxmlElement("w:gridCol"); col.set(qn("w:w"),str(width)); grid.append(col)
    for row in table.rows:
        trPr = row._tr.get_or_add_trPr()
        if trPr.find(qn("w:cantSplit")) is None:
            trPr.append(OxmlElement("w:cantSplit"))
        for i,cell in enumerate(row.cells):
            cell.width=Inches(widths[i]/1440); cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; margins(cell)
            tcW=cell._tc.get_or_add_tcPr().find(qn("w:tcW")); tcW.set(qn("w:w"),str(widths[i])); tcW.set(qn("w:type"),"dxa")

def add_table(headers, rows, widths):
    t=doc.add_table(rows=1, cols=len(headers)); t.style="Table Grid"; set_table_geometry(t,widths)
    for i,h in enumerate(headers): shade(t.rows[0].cells[i],PALE); p=t.rows[0].cells[i].paragraphs[0]; p.paragraph_format.space_after=Pt(0); font(p.add_run(h),10,bold=True,color=NAVY)
    for row in rows:
        cells=t.add_row().cells
        for i,value in enumerate(row): p=cells[i].paragraphs[0]; p.paragraph_format.space_after=Pt(0); font(p.add_run(str(value)),9.5,color="202733")
    set_table_geometry(t,widths); doc.add_paragraph().paragraph_format.space_after=Pt(0)
    return t

def add_step(title, owner, action, check=None):
    p=doc.add_paragraph(style="List Number"); font(p.add_run(title+" — "),11,bold=True,color=NAVY); font(p.add_run(action),11)
    q=doc.add_paragraph(); q.paragraph_format.left_indent=Inches(.375); q.paragraph_format.space_after=Pt(4); font(q.add_run(f"Owner: {owner}"),9.5,bold=True,color=GRAY)
    if check:
        q.add_run("   |   "); font(q.add_run(f"Evidence: {check}"),9.5,color=GRAY)

def callout(label, text, color=CYAN):
    t=doc.add_table(rows=1,cols=1); set_table_geometry(t,[9360]); shade(t.cell(0,0),PALE)
    p=t.cell(0,0).paragraphs[0]; p.style=note; font(p.add_run(label.upper()+": "),10.5,bold=True,color=BLUE); font(p.add_run(text),10.5,color=NAVY)
    doc.add_paragraph().paragraph_format.space_after=Pt(0)

def page_break(): doc.add_page_break()

# Running furniture
header=sec.header.paragraphs[0]; header.alignment=WD_ALIGN_PARAGRAPH.LEFT; font(header.add_run("DURANKI BANKING  |  STANDARD OPERATING PROCEDURE"),8.5,bold=True,color=GRAY)
footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.RIGHT; font(footer.add_run("Controlled document  |  Internal demo operations  |  Page "),8,color=GRAY)
fld=OxmlElement("w:fldSimple"); fld.set(qn("w:instr"),"PAGE"); footer._p.append(fld)

# Cover
if LOGO.exists():
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(70); p.add_run().add_picture(str(LOGO),width=Inches(1.55))
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.style="Title"; font(p.add_run("STANDARD OPERATING PROCEDURE"),30,bold=True,color=NAVY)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; font(p.add_run("Duranki Digital Banking & FLEXCUBE Integration Demonstration"),16,bold=True,color=BLUE)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(14); font(p.add_run("Customer Online Banking • Digital Onboarding • KYC Operations Portal"),11,color=GRAY)
doc.add_paragraph().paragraph_format.space_before=Pt(90)
add_table(["Document owner","Version","Effective date","Classification"],[["Duranki Platform Operations","1.2","1 September 2026","Bank Demonstration"]],[2600,1500,2200,3060])
callout("Purpose", "Provide a controlled, repeatable bank demonstration of the Duranki digital banking experience and the proposed integration with the bank's existing Oracle FLEXCUBE environment. This SOP covers customer banking, digital onboarding, KYC operations and the integration control model.")
page_break()

doc.add_heading("Document control",0)
add_table(["Field","Value"],[("Document ID","DUR-BANK-SOP-001"),("Process owner","Duranki Platform Operations"),("Approver","Head of Banking Operations / Product Owner"),("Review cycle","Quarterly and after every material release"),("Related systems","Duranki channels, Operations portal and the bank's Oracle FLEXCUBE environment"),("Access location","Use the approved demonstration address supplied by the system administrator")],[2700,6660])
doc.add_heading("Approval and revision history",1)
add_table(["Version","Date","Change","Author / approver"],[("1.2","1 Sep 2026","Bank demonstration and FLEXCUBE integration procedure","Duranki Platform Operations")],[1200,1600,4300,2260])
doc.add_heading("Contents",1)
for item in ["1. Purpose and scope","2. Roles and access","3. FLEXCUBE integration overview and prerequisites","4. Start-up procedure","5. Customer onboarding and KYC submission","6. KYC review and approval","7. Customer banking operations","8. QR merchant payments","9. Daily controls and monitoring","10. Backup and recovery","11. Shutdown procedure","12. Incident handling and troubleshooting","13. Security and compliance controls","14. Bank demonstration runbook","Appendix A: Integration discovery record","Appendix B: Operator checklists"]:
    doc.add_paragraph(item,style="List Bullet")

doc.add_heading("1. Purpose and scope",1)
doc.add_paragraph("This SOP applies to Duranki and bank personnel who prepare, demonstrate, assess or support the proposed digital banking solution. It explains the customer experience, the bank-controlled KYC workflow and the proposed exchange of approved customer, account, balance, payment and status information with Oracle FLEXCUBE.")
callout("Important", "The platform is a demonstration environment. No real money, production identity records or live payment credentials may be entered or processed.")

doc.add_heading("2. Roles and access",1)
add_table(["Role","Permitted activities","Key control"],[("Demo administrator","Prepare the application; verify availability; prepare demo data","Must complete pre-demo checklist"),("Bank compliance officer","Review customer profiles and approve or reject KYC","Decision and notes are retained"),("Bank FLEXCUBE owner","Confirm system-of-record rules, services and product mappings","Approves the target integration design"),("Integration lead","Explain data flows, controls, error handling and reconciliation","No direct core-database access"),("Customer / presenter","Register, sign in and demonstrate banking functions","Use demo-only identity and transaction data")],[1800,4500,3060])

doc.add_heading("3. FLEXCUBE integration overview and prerequisites",1)
doc.add_paragraph("Duranki provides the customer and bank-operations experience while Oracle FLEXCUBE remains the bank's core banking system of record for the products and records agreed during discovery. Integration is performed through bank-approved FLEXCUBE REST services or Integration Gateway web services. Duranki will not connect directly to the FLEXCUBE database.")
add_table(["Flow","Duranki responsibility","FLEXCUBE responsibility","Control"],[("Approved customer","Capture onboarding and release only after KYC approval","Create or update the agreed customer/party record","Idempotent request and bank reference"),("Accounts and balances","Request and present authorized customer data","Return eligible accounts and current balances","Read-only query with access checks"),("Payment instruction","Validate journey and submit approved instruction","Perform core validation, posting and status response","Unique reference; no silent retry"),("Transactions and status","Present reconciled results and customer status","Provide transaction, rejection and event status","Timestamped audit trail"),("Operational events","Consume agreed notifications or poll safely","Publish or expose agreed events","Monitoring and exception queue")],[1800,3000,3000,1560])
callout("Integration principle", "FLEXCUBE remains authoritative for core customer, account, balance and posting outcomes where agreed with the bank. Duranki orchestrates the digital journey and retains only the data required for channel operation, consent, evidence and audit.")
doc.add_paragraph("Prerequisites",style="Heading 2")
for text in ["The bank confirms its FLEXCUBE release and enabled integration services.","The approved demonstration or test environment and bank service accounts are available.","Customer, branch, currency, product and transaction-code mappings are agreed.","Authentication, encryption, network allow-listing, timeouts and audit requirements are documented.","Test cases cover success, rejection, duplicate, timeout and reconciliation outcomes.","Only demonstration data is used."]:
    doc.add_paragraph(text,style="List Bullet")

page_break()
doc.add_heading("4. Start-up procedure",1)
add_step("Confirm application availability","Demo administrator","Open the approved Duranki application address.","The login page appears without an error")
add_step("Confirm customer access","Demo administrator","Open the customer sign-in page and confirm that it responds.","Customer sign-in is available")
add_step("Confirm staff access","Demo administrator","Open the Operations portal and confirm that authorized staff can sign in.","The KYC review queue is available")
add_step("Confirm integration readiness","Integration lead","Confirm that the agreed FLEXCUBE test services or demonstration simulator are available.","The integration status check succeeds")
add_step("Load demonstration data","Demo administrator","Confirm that the approved sample customer, accounts and transactions are present.","Sample records display correctly")
add_step("Prepare presentation devices","Demo administrator","Confirm the browser, mobile layout and camera permission required for QR scanning.","Web, mobile and QR views are ready")
add_step("Perform smoke checks","Demo administrator","Confirm customer login, staff login, account view and KYC list load without errors.","Pre-demo checklist signed")
callout("Stop condition", "Do not continue if customer or staff sign-in is unavailable, the KYC queue does not load, the FLEXCUBE test connection is unavailable, banking records appear inconsistent, or the application remains on a loading screen for more than 15 seconds.")

doc.add_heading("5. Customer onboarding and KYC submission",1)
add_step("Open registration","Customer","Select Register as a new client from the login page.")
add_step("Capture secure sign-in details","Customer","Enter full legal name, email, mobile number and a password of at least 10 characters.")
add_step("Capture identity details","Customer","Enter a demo ID/passport number and date of birth. Never use a real identity document in the demonstration environment.")
add_step("Capture address and financial profile","Customer","Complete address, occupation, source of funds, tax-residency and politically exposed person declarations.")
add_step("Review and consent","Customer","Confirm that the details are accurate and submit the profile.")
add_step("Confirm pending status","Customer / operator","Verify that online banking remains locked and the KYC reference is displayed.","KYC record appears in Operations portal")

doc.add_heading("6. KYC review and approval",1)
add_step("Open Bank Staff","Compliance officer","Select Bank Staff and sign in with authorized demo staff credentials.")
add_step("Locate the applicant","Compliance officer","Use the Pending filter or search by name, email or ID.")
add_step("Review the customer information file","Compliance officer","Check identity, contact, address, occupation, source of funds, tax residency and PEP status.")
add_step("Make a decision","Compliance officer","Approve only when required demo checks are complete; otherwise reject/request changes and record a clear note.")
add_step("Release the approved record","Bank compliance officer / integration service","After approval, submit the agreed customer creation or update request to FLEXCUBE and retain the returned bank reference.","Approved record is linked to the bank reference")
add_step("Verify access outcome","Compliance officer / customer","Sign out of staff access. The customer signs in again; approved customers reach the dashboard, rejected customers see remediation guidance.")
callout("Four-eyes control", "For formal stakeholder demonstrations, the presenter should not approve their own KYC record. Use a separate compliance-officer role.")

doc.add_heading("7. Customer banking operations",1)
add_table(["Operation","Procedure","Expected result"],[("View balances","Open Home or Accounts.","Everyday, savings and linked-account balances display."),("Send money / pay bill","Enter recipient, amount and reference; submit.","Balance decreases and a transaction is recorded."),("Link external account","Choose Add account; select mobile money, bank, wallet or investment; save.","Account appears in the combined portfolio."),("Remove linked account","Select the remove control on an external account.","Only the selected linked account is deleted."),("Freeze card","Open Cards and toggle card status.","Everyday account status changes to frozen/active."),("Edit profile","Open Profile, update details and save.","Customer profile changes are retained."),("Sign out","Select Sign out.","The session ends and the login page appears.")],[1900,4200,3260])
callout("Transaction safety", "The demonstration rejects payments with an invalid account, non-positive amount or insufficient balance. Never attempt to bypass these controls during a demonstration.")

doc.add_heading("8. QR merchant payments",1)
add_step("Create merchant request","Merchant / presenter","Open Payments, select Create merchant QR, enter merchant, amount and reference, then generate the request.")
add_step("Present the code","Merchant / presenter","Display the generated QR/payment code. It expires after 15 minutes.")
add_step("Scan or enter code","Customer","Open Scan to pay, use the camera or enter the payment code manually.")
add_step("Review before approval","Customer","Verify merchant, amount and reference. Cancel if any detail is incorrect.")
add_step("Confirm payment","Customer","Approve the payment.","QR status becomes paid; balance and transaction history update")

doc.add_heading("9. Daily controls and monitoring",1)
add_table(["Control","Frequency","Owner","Evidence"],[("Application availability check","Before each session","Demo administrator","Customer and staff pages observed"),("FLEXCUBE connection check","Before each session","Integration lead","Approved test service responds"),("Pending KYC queue review","Daily during active testing","Compliance officer","Queue contains no unexplained records"),("Interface exception review","After demonstrations","Integration lead","Rejected and timed-out messages reviewed"),("Demo-data validation","Before external demos","Product owner","No real PII or payment data present"),("Reconciliation check","After transactional demos","Bank operations","Duranki and FLEXCUBE references agree")],[3100,1700,1900,2860])

doc.add_heading("10. Backup and recovery",1)
doc.add_paragraph("Preserve demonstration data using the approved application backup process and an authorized encrypted location. Backups must not contain real customer data.")
add_step("Create backup","Demo administrator","Run the approved application-data backup procedure to an authorized encrypted location.","Backup completes successfully")
add_step("Record metadata","Demo administrator","Record date, environment, schema version and operator.")
add_step("Test restore","Application support / administrator","Restore into an isolated test environment and complete the application checks.","Login and KYC records are accessible")
add_step("Retain and dispose","Process owner","Follow the agreed demo retention period; securely remove expired backups.")

doc.add_heading("11. Shutdown procedure",1)
add_step("Finish active workflows","Presenter","Complete or cancel payments and KYC decisions. Sign out customer and staff sessions.")
add_step("End active presentation","Demo administrator","Confirm that no demonstration activity is still in progress and close the presentation view.")
add_step("Close the application","Demo administrator","Close the customer and staff application sessions after all users have signed out.","No demonstration sessions remain active")
add_step("Preserve records when required","Demo administrator","Retain required demonstration records and backups according to the approved retention period.")
add_step("Confirm completion","Demo administrator","Complete the post-demonstration checklist and record any incident or stakeholder feedback.")

doc.add_heading("12. Incident handling and troubleshooting",1)
add_table(["Symptom","Likely cause","Action"],[("Application cannot be reached","Application or network unavailable","Confirm the approved address and network connection; contact application support."),("FLEXCUBE request rejected","Mapping, authorization or core validation","Display the controlled message; retain references and review with the bank owner."),("FLEXCUBE request times out","Network or service latency","Do not submit blindly; query status using the original reference and reconcile."),("Banking information does not load","Integration service unavailable or access expired","Sign in again; if unresolved, preserve the error details and escalate."),("Login rejected","Incorrect credentials or account status","Confirm demo credentials; ask authorized staff to review the user and KYC status."),("KYC list empty","Staff session, filter or application issue","Sign in again, clear filters and confirm the correct staff role; escalate if unresolved."),("Unexpected transaction result","Validation or integration error","Stop further transactions, preserve both system references and notify the owners.")],[2300,2800,4260])
callout("Escalation", "Stop the demonstration and notify the product owner if data appears inconsistent, an unauthorized user gains access, or any transaction behaves unexpectedly. Preserve logs and do not delete evidence.")

doc.add_heading("13. Security and compliance controls",1)
for text in ["Use encrypted connections, secure sessions, strong secrets and approved access restrictions.","Replace all demonstration credentials before deployment; never share production secrets.","Apply least privilege and maintain separate customer, staff and administrative roles.","Record KYC decisions and banking changes in immutable audit logs.","Encrypt backups and data in transit; establish retention and deletion schedules.","Use an approved identity-verification provider for production KYC.","Conduct vulnerability scanning, penetration testing and security review before launch.","Treat this application as a prototype until legal, compliance, security and banking-product approvals are complete."]:
    doc.add_paragraph(text,style="List Bullet")

doc.add_heading("14. Bank demonstration runbook",1)
add_table(["Time","Presenter action","Audience outcome"],[("0-4 min","Confirm bank objectives, FLEXCUBE context and demo disclaimer.","Agrees scope and system-of-record principles."),("4-9 min","Register a new customer and submit KYC.","Sees seamless digital onboarding."),("9-14 min","Review/approve KYC and show the proposed FLEXCUBE handoff.","Sees bank-controlled authorization and reference linkage."),("14-21 min","Sign in as approved customer; show FLEXCUBE-sourced account and balance views.","Sees the target online banking experience."),("21-26 min","Demonstrate a payment, status response, reconciliation and exception path.","Sees controlled core-banking integration."),("26-30 min","Discuss security, data residency, delivery phases and bank decisions.","Understands next steps and responsibilities.")],[1300,4300,3760])

doc.add_heading("Appendix A: Integration discovery record",1)
add_table(["Decision area","To be confirmed with the bank"],[("FLEXCUBE baseline","Release, modules, deployment model and enabled service catalogue"),("Interface pattern","REST services, Integration Gateway web services, outbound notifications or approved combination"),("System of record","Ownership for party/customer, KYC, accounts, balances, payments and transaction status"),("Mappings","Branches, currencies, products, account classes, transaction codes and status codes"),("Security","Identity, service accounts, certificates, encryption, network zones and access approvals"),("Operational controls","Timeouts, retry rules, idempotency, reconciliation, monitoring and support ownership"),("Non-functional targets","Availability, response times, throughput, recovery objectives, retention and data residency"),("Delivery path","Discovery, sandbox integration, system testing, user acceptance, pilot and production rollout")],[3000,6360])
doc.add_heading("Demonstration access and controls",2)
for text in ["Use the approved Duranki demonstration address and authorized bank-staff account.","Use a FLEXCUBE non-production environment or approved simulator; never use live customer data.","Do not place passwords or service credentials in this SOP; share them through an approved secure channel.","Label simulated integration responses clearly when a bank test connection is not available.","Retain a Duranki reference and a FLEXCUBE or simulator reference for every demonstrated handoff.","Record bank questions, required mappings and decisions in the integration discovery record."]:
    doc.add_paragraph(text,style="List Bullet")

page_break()
p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(12)
doc.add_heading("Appendix B: Operator checklists",1)
doc.add_heading("Pre-demonstration checklist",2)
for text in ["☐ Approved Duranki application address opens","☐ Customer and staff sign-in pages are available","☐ FLEXCUBE test service or clearly labelled simulator is available","☐ Bank mappings and demonstration scenarios are agreed","☐ KYC review queue and customer dashboard load","☐ Staff and approved demo customer login succeed","☐ No real customer data is present","☐ QR/camera permissions are ready or manual-code fallback is planned"]:
    doc.add_paragraph(text)
doc.add_heading("Post-demonstration checklist",2)
for text in ["☐ Customer and staff sessions signed out","☐ Duranki and FLEXCUBE/simulator references reconciled","☐ Interface exceptions and rejected messages reviewed","☐ Demo-only records reviewed and required backup completed","☐ Browser and mobile sessions closed","☐ Bank decisions and integration questions recorded","☐ Incidents or stakeholder feedback recorded"]:
    doc.add_paragraph(text)

doc.add_heading("Operator acknowledgement",2)
add_table(["Name","Role","Date","Signature"],["","","",""] if False else [("","","","")],[2500,2200,1800,2860])

doc.core_properties.title = "Duranki Banking App Standard Operating Procedure"
doc.core_properties.subject = "Operation of the Duranki digital banking demonstration platform"
doc.core_properties.author = "Duranki Platform Operations"
doc.core_properties.keywords = "Duranki, banking, SOP, KYC, digital onboarding, operations"
doc.save(OUT)
print(OUT)
