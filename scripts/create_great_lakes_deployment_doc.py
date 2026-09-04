from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

ROOT = Path(r"C:\Users\johan\Documents\Codex\2026-08-02\c")
OUT = ROOT / "outputs" / "Great Lakes Deployment.docx"
LOGO = ROOT / "public" / "great-lakes-bank-logo.png"

NAVY = "063A74"
DEEP = "092D50"
TEAL = "0AA8BA"
GOLD = "E9B53F"
LIGHT = "EEF5F8"
PALE_GOLD = "FFF7E2"
RED = "9B1C1C"
GRAY = "63798A"

doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(8.5), Inches(11)
sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(0.8)
sec.header_distance = sec.footer_distance = Inches(0.35)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Arial"; normal.font.size = Pt(10.5); normal.font.color.rgb = RGBColor.from_string(DEEP)
normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.15
for name, size, before, after, color in [("Title",30,0,10,NAVY),("Subtitle",13,0,16,GRAY),("Heading 1",18,18,8,NAVY),("Heading 2",14,14,6,TEAL),("Heading 3",11.5,10,4,DEEP)]:
    st = styles[name]; st.font.name="Arial"; st.font.size=Pt(size); st.font.bold=name!="Subtitle"; st.font.color.rgb=RGBColor.from_string(color)
    st.paragraph_format.space_before=Pt(before); st.paragraph_format.space_after=Pt(after); st.paragraph_format.keep_with_next=True
for name in ["List Bullet","List Number"]:
    st=styles[name]; st.font.name="Arial"; st.font.size=Pt(10.5); st.paragraph_format.left_indent=Inches(.38); st.paragraph_format.first_line_indent=Inches(-.19); st.paragraph_format.space_after=Pt(4); st.paragraph_format.line_spacing=1.15

def shade(cell, fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=tcPr.find(qn("w:shd"))
    if shd is None: shd=OxmlElement("w:shd"); tcPr.append(shd)
    shd.set(qn("w:fill"),fill)

def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc=cell._tc; tcPr=tc.get_or_add_tcPr(); tcMar=tcPr.first_child_found_in("w:tcMar")
    if tcMar is None: tcMar=OxmlElement("w:tcMar"); tcPr.append(tcMar)
    for m,v in (("top",top),("start",start),("bottom",bottom),("end",end)):
        node=tcMar.find(qn(f"w:{m}"))
        if node is None: node=OxmlElement(f"w:{m}"); tcMar.append(node)
        node.set(qn("w:w"),str(v)); node.set(qn("w:type"),"dxa")

def page_break(): doc.add_page_break()

def p(text="", bold=False, italic=False, color=None, size=None, align=None, after=6):
    para=doc.add_paragraph(); para.paragraph_format.space_after=Pt(after)
    if align is not None: para.alignment=align
    r=para.add_run(text); r.bold=bold; r.italic=italic; r.font.name="Arial"
    if color: r.font.color.rgb=RGBColor.from_string(color)
    if size: r.font.size=Pt(size)
    return para

def bullet(text): doc.add_paragraph(text, style="List Bullet")
def number(text): doc.add_paragraph(text, style="List Number")
def h1(text): doc.add_heading(text, level=1)
def h2(text): doc.add_heading(text, level=2)
def h3(text): doc.add_heading(text, level=3)

def callout(label, text, fill=LIGHT, label_color=TEAL):
    t=doc.add_table(rows=1, cols=1); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False; t.columns[0].width=Inches(6.7)
    c=t.cell(0,0); c.width=Inches(6.7); shade(c,fill); set_cell_margins(c,160,180,160,180)
    pr=c.paragraphs[0]; pr.paragraph_format.space_after=Pt(3)
    r=pr.add_run(label.upper()); r.bold=True; r.font.name="Arial"; r.font.size=Pt(9); r.font.color.rgb=RGBColor.from_string(label_color)
    pr2=c.add_paragraph(text); pr2.paragraph_format.space_after=Pt(0); pr2.paragraph_format.line_spacing=1.1
    doc.add_paragraph().paragraph_format.space_after=Pt(2)

def code(lines):
    t=doc.add_table(rows=1,cols=1); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False; c=t.cell(0,0); shade(c,"10283F"); set_cell_margins(c,140,160,140,160)
    pr=c.paragraphs[0]; pr.paragraph_format.space_after=Pt(0)
    for i,line in enumerate(lines.strip().splitlines()):
        if i: pr.add_run("\n")
        r=pr.add_run(line); r.font.name="Consolas"; r.font.size=Pt(8.7); r.font.color.rgb=RGBColor(245,248,250)
    doc.add_paragraph().paragraph_format.space_after=Pt(1)

def checklist(items):
    for item in items:
        pr=doc.add_paragraph(); pr.paragraph_format.left_indent=Inches(.12); pr.paragraph_format.space_after=Pt(4)
        r=pr.add_run("☐  "); r.font.name="Arial"; r.font.color.rgb=RGBColor.from_string(TEAL)
        pr.add_run(item)

def add_table(headers, rows, widths):
    t=doc.add_table(rows=1,cols=len(headers)); t.autofit=False; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,(head,w) in enumerate(zip(headers,widths)):
        c=t.rows[0].cells[i]; c.width=Inches(w); shade(c,NAVY); set_cell_margins(c)
        r=c.paragraphs[0].add_run(head); r.bold=True; r.font.name="Arial"; r.font.size=Pt(9); r.font.color.rgb=RGBColor(255,255,255)
    for ridx,row in enumerate(rows):
        cells=t.add_row().cells
        for i,(value,w) in enumerate(zip(row,widths)):
            cells[i].width=Inches(w); set_cell_margins(cells[i]); cells[i].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if ridx%2: shade(cells[i],"F7FAFC")
            cells[i].text=str(value)
            for run in cells[i].paragraphs[0].runs: run.font.name="Arial"; run.font.size=Pt(9)
    return t

# Running furniture
header=sec.header.paragraphs[0]; header.text="GREAT LAKES BANK  |  DEPLOYMENT GUIDE"; header.alignment=WD_ALIGN_PARAGRAPH.RIGHT
for r in header.runs: r.font.name="Arial"; r.font.size=Pt(8); r.font.bold=True; r.font.color.rgb=RGBColor.from_string(NAVY)
footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER
rr=footer.add_run("Controlled demonstration environment • Great Lakes Bank"); rr.font.name="Arial"; rr.font.size=Pt(8); rr.font.color.rgb=RGBColor.from_string(GRAY)

# Cover
if LOGO.exists():
    para=doc.add_paragraph(); para.alignment=WD_ALIGN_PARAGRAPH.CENTER; para.paragraph_format.space_after=Pt(24); para.add_run().add_picture(str(LOGO),width=Inches(1.55))
p("DEPLOYMENT & UPDATE GUIDE",bold=True,color=GOLD,size=10,align=WD_ALIGN_PARAGRAPH.CENTER,after=8)
title=doc.add_paragraph(style="Title"); title.alignment=WD_ALIGN_PARAGRAPH.CENTER; title.add_run("Great Lakes Deployment")
p("A beginner-friendly guide to publishing the digital banking demonstration on Evans's server and maintaining it through Git.",color=GRAY,size=13,align=WD_ALIGN_PARAGRAPH.CENTER,after=24)
callout("Purpose","Use this guide for a controlled demonstration with synthetic data. The current mock FLEXCUBE, SMS and biometric providers must not be represented as live bank integrations.",PALE_GOLD,GOLD)
p("Prepared: 4 September 2026",bold=True,color=NAVY,size=10,align=WD_ALIGN_PARAGRAPH.CENTER,after=4)
p("Project location on the development computer",color=GRAY,size=9,align=WD_ALIGN_PARAGRAPH.CENTER,after=2)
p(r"C:\Users\johan\Documents\Codex\2026-08-02\c",bold=True,color=TEAL,size=9,align=WD_ALIGN_PARAGRAPH.CENTER)

page_break(); h1("1. Understand what will be published")
p("The recommended design uses one public subdomain, for example banking.evansdomain.com. Visitors see one address while Nginx privately forwards requests to the correct service.")
add_table(["Public request","Internal destination","Purpose"],[
    ("https://banking.evansdomain.com/","127.0.0.1:3002","Customer banking and operations portal"),
    ("https://banking.evansdomain.com/api/...","127.0.0.1:3001","Express digital banking API"),
    ("MariaDB","Docker private network","Persistent banking demonstration data"),
],[2.55,2.0,2.15])
h2("What is already in the repository")
bullet("Frontend: responsive customer banking and bank operations portal.")
bullet("Backend: Express and TypeScript modular banking API.")
bullet("Database: MariaDB schema and persistent Docker volumes.")
bullet("Integration abstractions: mock FLEXCUBE, SMS and biometrics for demonstration use.")
callout("Important","The frontend is not currently included in the Docker Compose stack. Docker runs MariaDB and the API; Node.js runs the frontend as a separate managed service.")
h2("Information to obtain from Evans")
checklist(["Server public IP address","SSH username and SSH key or temporary password","Server operating system and version","Permission to use sudo","Chosen subdomain","Access to the domain DNS records","Confirmation that inbound ports 80 and 443 are open"])

h1("2. Put the project in a private Git repository")
p("Git becomes the master copy of the application. The server downloads the first version with git clone and later updates with git pull.")
h2("On your Windows development computer")
number("Open PowerShell in the project folder.")
code(r'''cd "C:\Users\johan\Documents\Codex\2026-08-02\c"''')
number("Confirm the repository status and review what will be uploaded.")
code("git status")
number("Ensure secrets and generated files are excluded through .gitignore. Never commit .env files, passwords, database exports, node_modules, build output or uploaded KYC documents.")
number("Add and commit the approved source changes.")
code('''git add .
git commit -m "Prepare Great Lakes Bank deployment"''')
number("Connect the private remote repository. Run the remote-add command only once.")
code('''git branch -M main
git remote add origin https://github.com/YOUR-ACCOUNT/YOUR-PRIVATE-REPOSITORY.git
git push -u origin main''')
callout("Private repository","Banking source code and deployment configuration should not be placed in a public repository. Give Evans or the deployment engineer only the access required to clone it.")

page_break(); h1("3. Prepare the server")
p("These commands assume Evans's server uses a supported Ubuntu Linux release. For Windows Server, cPanel, Plesk or another platform, stop and obtain platform-specific instructions.")
h2("Connect from Windows")
code("ssh USERNAME@EVANS_SERVER_IP")
h2("Install the runtime")
p("Install Docker Engine and the Docker Compose plugin from Docker's official Ubuntu repository. Install Node.js 22 or newer, Git, Nginx and Certbot using the server administrator's approved package sources.")
code('''docker --version
docker compose version
node --version
npm --version
git --version
nginx -v''')
p("Node must be version 22 or newer. Docker Compose should use the modern docker compose command.")
h2("Create the application location")
code('''sudo mkdir -p /opt/great-lakes-bank
sudo chown "$USER":"$USER" /opt/great-lakes-bank''')
h2("Clone the private repository")
code('''cd /opt
git clone https://github.com/YOUR-ACCOUNT/YOUR-PRIVATE-REPOSITORY.git great-lakes-bank
cd /opt/great-lakes-bank
git status''')
callout("Authentication","Use a deploy key or narrowly scoped repository credential. Do not paste a personal GitHub password into scripts or configuration files.")

h1("4. Configure DNS")
number("Open the DNS management screen for Evans's domain.")
number("Create an A record named banking pointing to Evans's server IP address.")
add_table(["Field","Value"],[('Type','A'),('Name','banking'),('Value','EVANS_SERVER_IP'),('TTL','Automatic or provider default')],[1.75,4.95])
number("Wait for DNS to resolve, then verify it from your computer.")
code("nslookup banking.evansdomain.com")

h1("5. Create deployment-only configuration")
p("The repository contains local demonstration values. Evans or the deployment engineer must create private server configuration that is never committed to Git.")
h2("Required public addresses")
code('''NODE_ENV=production
APP_BANK_NAME=Great Lakes Bank
FRONTEND_ORIGIN=https://banking.evansdomain.com
NEXT_PUBLIC_API_ORIGIN=https://banking.evansdomain.com
NEXT_PUBLIC_APP_BANK_NAME=Great Lakes Bank
NEXT_PUBLIC_DEMO_MODE=true''')
h2("Demonstration integration selection")
code('''DEMO_MODE=true
CORE_BANKING_PROVIDER=mock
FACE_PROVIDER=mock
LIVENESS_PROVIDER=mock
FINGERPRINT_PROVIDER=mock
SMS_PROVIDER=mock''')
h2("Secrets that must be replaced")
bullet("MariaDB application password and root password")
bullet("JWT and access-token signing secrets")
bullet("OTP secret")
bullet("Administrator, KYC officer and operations passwords")
bullet("Any future FLEXCUBE, SMS, biometric or object-storage credentials")
p("Generate a different strong secret for every secret field:")
code("openssl rand -hex 48")
callout("Never commit secrets","Do not store production or server passwords in docker-compose.yml, source files, Git commits, screenshots, email or messaging conversations.","FDECEC",RED)

h1("6. Keep internal services private")
p("The public internet must not connect directly to MariaDB, the API port or the frontend port. Nginx should be the only public entry point.")
bullet("Do not publish MariaDB port 3306 publicly. Prefer no MariaDB ports entry at all.")
bullet("Bind the API only to 127.0.0.1:3001.")
bullet("Bind the frontend only to 127.0.0.1:3002.")
bullet("Allow public inbound traffic only on ports 80 and 443, plus restricted SSH administration.")
callout("Docker firewall warning","Docker-published ports can interact with host firewall rules. Evans's administrator must verify the effective Docker and host firewall configuration, not only the visible UFW rules.",PALE_GOLD,GOLD)

h1("7. Start MariaDB and the backend API")
code('''cd /opt/great-lakes-bank
docker compose -f backends/docker-compose.yml up --build -d''')
h2("Check the service state")
code('''docker compose -f backends/docker-compose.yml ps
docker compose -f backends/docker-compose.yml logs --tail=100 api
curl http://127.0.0.1:3001/health/ready''')
p("A healthy readiness response should report the application and database online. Mock providers may appear as MOCK on administration screens.")
h2("Useful troubleshooting")
code('''docker compose -f backends/docker-compose.yml logs --tail=200 api
docker compose -f backends/docker-compose.yml logs --tail=200 mariadb''')
callout("Protect the data","Never run docker compose down -v on the deployed server. The -v option removes the database and document-storage volumes.","FDECEC",RED)

h1("8. Build and start the frontend")
p("The NEXT_PUBLIC values must be set before the build because the public API address is compiled into the frontend output.")
code('''cd /opt/great-lakes-bank
npm ci
npm run build
npm run start -- --host 127.0.0.1 --port 3002''')
p("From a second SSH session, verify:")
code("curl http://127.0.0.1:3002")
h2("Make it survive logout and restart")
p("A server administrator must create a restricted systemd service named great-lakes-bank-frontend. It should use /opt/great-lakes-bank as its working directory, run the npm start command above, restart on failure and start after networking is ready.")
code('''sudo systemctl daemon-reload
sudo systemctl enable --now great-lakes-bank-frontend
sudo systemctl status great-lakes-bank-frontend''')

h1("9. Configure Nginx")
p("Create an Nginx site for the chosen subdomain. Replace the example domain before enabling it.")
code('''server {
    listen 80;
    server_name banking.evansdomain.com;

    location /api/ {
        proxy_pass http://127.0.0.1:3001;
        proxy_set_header Host $host;
        proxy_set_header X-Real-IP $remote_addr;
        proxy_set_header X-Forwarded-For $proxy_add_x_forwarded_for;
        proxy_set_header X-Forwarded-Proto $scheme;
    }

    location /health {
        proxy_pass http://127.0.0.1:3001;
        proxy_set_header Host $host;
        proxy_set_header X-Forwarded-Proto $scheme;
    }

    location / {
        proxy_pass http://127.0.0.1:3002;
        proxy_set_header Host $host;
        proxy_set_header X-Real-IP $remote_addr;
        proxy_set_header X-Forwarded-For $proxy_add_x_forwarded_for;
        proxy_set_header X-Forwarded-Proto $scheme;
    }
}''')
h2("Validate before applying")
code('''sudo nginx -t
sudo systemctl reload nginx''')
callout("Stop on failure","If nginx -t reports an error, do not reload Nginx. Copy the exact error for the server administrator to correct.")

h1("10. Enable HTTPS")
p("The banking interface must not be presented publicly over plain HTTP.")
code('''sudo certbot --nginx -d banking.evansdomain.com
sudo certbot renew --dry-run''')
p("Open https://banking.evansdomain.com and confirm the browser shows a valid secure connection.")

h1("11. First-deployment acceptance test")
checklist([
    "Great Lakes Bank login page loads over HTTPS.",
    "/health/ready reports application and database online.",
    "Customer login and logout work.",
    "KYC officer login works.",
    "Existing-customer activation reaches OTP and KYC onboarding.",
    "Synthetic document and selfie uploads work.",
    "KYC appears in the operations queue and can be reviewed.",
    "Approval changes the customer's access state.",
    "Approved customer accounts and balances load through the backend.",
    "Beneficiary creation and internal demonstration transfer work.",
    "Transfer and KYC events appear in Operations and Audit.",
    "The responsive interface works from a phone using mobile data.",
    "After a controlled server restart, Docker and the frontend return automatically.",
])
callout("Use synthetic data only","Do not upload real passports, identity documents, customer photos, account numbers or personal information into the demonstration environment.",PALE_GOLD,GOLD)

h1("12. Routine Git update procedure")
p("Use this procedure whenever an approved version has been pushed to the main branch.")
h2("A. On the development computer")
code('''git status
git add .
git commit -m "Describe the approved change"
git push origin main''')
h2("B. On Evans's server")
number("Back up the database and uploaded documents before deploying.")
number("Download the approved source version.")
code('''cd /opt/great-lakes-bank
git status
git pull --ff-only origin main''')
number("Install exact frontend dependencies and run tests/build.")
code('''npm ci
npm test
npm run build''')
number("Run backend tests before rebuilding the containers.")
code('''cd /opt/great-lakes-bank/backends/api
npm ci
npm test
cd /opt/great-lakes-bank''')
number("Rebuild and restart the API, then restart the frontend.")
code('''docker compose -f backends/docker-compose.yml up --build -d
sudo systemctl restart great-lakes-bank-frontend''')
number("Verify the deployment.")
code('''curl https://banking.evansdomain.com/health/ready
docker compose -f backends/docker-compose.yml ps
sudo systemctl status great-lakes-bank-frontend''')

page_break(); h1("13. Safe rollback procedure")
p("Record the commit ID before every update. If the new version fails, return to the last approved commit without deleting database volumes.")
h2("Before updating")
code("git rev-parse HEAD")
h2("If rollback is required")
code('''cd /opt/great-lakes-bank
git checkout LAST_KNOWN_GOOD_COMMIT
npm ci
npm run build
docker compose -f backends/docker-compose.yml up --build -d
sudo systemctl restart great-lakes-bank-frontend''')
p("After the incident, create a corrective commit in the main repository. Do not continue making untracked edits directly on the server.")
callout("Database compatibility","A source rollback does not automatically reverse database migrations. Before deploying schema changes, the developer must provide a tested migration and rollback plan.",PALE_GOLD,GOLD)

h1("14. Backups and operations")
bullet("Create daily encrypted MariaDB backups.")
bullet("Back up KYC document storage separately.")
bullet("Keep at least one protected backup outside Evans's server.")
bullet("Define retention, access and deletion rules with Great Lakes Bank.")
bullet("Test restoration periodically; an untested backup is not a recovery plan.")
bullet("Monitor disk space, container health, certificate renewal and failed logins.")
bullet("Apply operating system and Docker security updates through a controlled maintenance process.")

h1("15. Responsibility checklist")
add_table(["Responsibility","Owner to confirm"],[
    ("Domain and DNS access","Evans / domain administrator"),
    ("Server access, firewall and operating system","Evans / server administrator"),
    ("Private Git repository and releases","Development owner"),
    ("Secrets and credential rotation","Authorised deployment administrator"),
    ("Database and document backups","Server/database administrator"),
    ("Demo script and synthetic data","Presentation owner"),
    ("Production security and compliance approval","Great Lakes Bank"),
],[4.55,2.15])

page_break(); h1("16. Quick command reference")
h2("Check services")
code('''docker compose -f backends/docker-compose.yml ps
sudo systemctl status great-lakes-bank-frontend
sudo systemctl status nginx''')
h2("View logs")
code('''docker compose -f backends/docker-compose.yml logs --tail=100 api
journalctl -u great-lakes-bank-frontend -n 100 --no-pager
sudo tail -n 100 /var/log/nginx/error.log''')
h2("Restart safely")
code('''docker compose -f backends/docker-compose.yml restart api
sudo systemctl restart great-lakes-bank-frontend
sudo systemctl reload nginx''')
h2("Update from Git")
code('''cd /opt/great-lakes-bank
git pull --ff-only origin main
npm ci && npm test && npm run build
docker compose -f backends/docker-compose.yml up --build -d
sudo systemctl restart great-lakes-bank-frontend''')

h1("17. Sources and official guidance")
for title,url in [
    ("Docker Engine installation for Ubuntu","https://docs.docker.com/engine/install/ubuntu/"),
    ("Docker Compose plugin installation","https://docs.docker.com/compose/install/linux/"),
    ("Certbot installation guidance","https://certbot.eff.org/instructions"),
    ("Certbot user guide and Nginx plugin","https://eff-certbot.readthedocs.io/en/stable/using.html"),
]:
    pr=doc.add_paragraph(style="List Bullet"); pr.add_run(title+": ").bold=True; pr.add_run(url)

callout("Final status","This guide publishes the current controlled demonstration architecture. Real customer onboarding or banking requires completed provider integrations, infrastructure hardening, independent penetration testing, operational controls, legal review and Great Lakes Bank approval.","FDECEC",RED)

OUT.parent.mkdir(parents=True,exist_ok=True)
doc.core_properties.title="Great Lakes Deployment"
doc.core_properties.subject="Git-based deployment and update guide for the Great Lakes Bank digital banking demonstration"
doc.core_properties.author="Great Lakes Bank Digital Banking Project"
doc.save(OUT)
print(OUT)
